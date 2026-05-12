"""Anthropic API 키 발급 가이드 v2 — 실제 한국어 UI 스크린샷 기반."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

DOC_DIR = Path(__file__).parent
IMG_DIR = DOC_DIR / "real"
OUT = DOC_DIR / "Anthropic_API_키_발급_가이드.docx"

doc = Document()
for s in doc.sections:
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)


def heading(text, level=1, color=(40, 80, 160)):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(20 - level * 2)
    run.font.color.rgb = RGBColor(*color)
    run.bold = True


def para(text="", size=11, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def callout(text, kind="info"):
    icons = {"info": "💡", "warn": "⚠", "note": "📌"}
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icons.get(kind, '•')}  {text}")
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10)
    color = {"info": (40, 80, 160), "warn": (180, 80, 30), "note": (60, 60, 60)}
    run.font.color.rgb = RGBColor(*color.get(kind, (60, 60, 60)))
    run.bold = True


def image(name, width_cm=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(IMG_DIR / name), width=Cm(width_cm))


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.italic = True


# ════════════ 표지 ════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Anthropic API 키 발급 가이드")
run.font.name = "맑은 고딕"
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(40, 80, 160)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("K-에듀파인 비전자문서등록 자동입력기 사용을 위한 사전 준비")
run.font.name = "맑은 고딕"
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 110, 130)

doc.add_paragraph()

# ── 개요 ──
heading("개요", 1)
para(
    "비전자문서등록 자동입력기는 PDF 의 내용을 자동으로 추출하기 위해 "
    "Anthropic Claude API 를 사용합니다. 본 가이드는 Claude Console 가입부터 "
    "결제 등록, API 키 발급, 자동입력기에 키를 등록하는 전 과정을 "
    "실제 화면 스크린샷과 함께 설명합니다."
)

callout("예상 소요 시간: 약 5-10분 (결제수단 + 카드 정보 입력 포함)", "info")
callout("최소 결제 금액: $5 (약 7,000원). PDF 1건당 약 10원 (Sonnet 4.5 기준) — 약 500건 사용 가능.", "info")
callout("이 키는 본인만 사용. 메신저·이메일·외부 사이트에 절대 공유 금지.", "warn")

doc.add_page_break()

# ════════════ 1단계 ════════════
heading("1단계 — Claude Console 가입", 1)
para("브라우저에서 다음 주소로 이동하세요:", 11)
para("https://console.anthropic.com/", 12, bold=True, color=(40, 80, 160))
para()
para("가입 방법:", bold=True)
para("  ● 빨간 박스 표시된 [Google 계정으로 계속하기] 클릭 (가장 빠름)")
para("  ● 또는 [이메일 입력] 칸에 이메일 입력 후 [이메일로 계속하기] → 이메일 인증")

image("1.png", 16)
caption("그림 1. Claude Console 가입/로그인 페이지")

callout("Google 계정 권장 — 별도 비밀번호 / 이메일 인증 불필요.", "info")

doc.add_page_break()

# ════════════ 2단계 ════════════
heading("2단계 — 크레딧 메뉴로 이동", 1)
para(
    "로그인 후 메인 대시보드가 나타납니다. "
    "좌측 사이드바 ★맨 아래★ 의 [크레딧] 메뉴를 클릭합니다."
)
para()
para("화면에서 확인:", bold=True)
para("  ● 좌측 사이드바 하단 빨간 박스의 [크레딧 USD 0.00] 클릭")
para("  ● (또는 우상단 [API 키 받기] 버튼으로도 진입 가능)")

image("2.png", 16)
caption("그림 2. 메인 대시보드 — 좌측 사이드바 [크레딧] 클릭")

callout("그림에서 [크레딧 USD 13.27] 로 표시된 곳은 이미 충전된 상태입니다. 처음엔 USD 0.00 으로 표시됩니다.", "note")

doc.add_page_break()

# ════════════ 3단계 ════════════
heading("3단계 — 크레딧 구매", 1)
para(
    "크레딧 잔액 페이지로 이동합니다. 처음에는 잔액이 $0 이고 결제수단도 미등록 상태입니다."
)
para()
para("순서:", bold=True)
para("  ① 페이지에서 빨간 박스 [크레딧 구매] 버튼 클릭")
para("  ② Stripe 결제 페이지로 이동 — 카드 정보 입력 + 충전 금액 선택 (최소 $5)")
para("  ③ 결제 완료 후 잔액이 표시됨 (예: US$13.27)")

image("3.png", 16)
caption("그림 3. 크레딧 페이지 — [크레딧 구매] 클릭")

callout("Auto-reload (자동 재충전) 설정도 가능 — 잔액이 일정 수준 이하로 떨어지면 자동 충전. 처음엔 비활성화된 상태.", "info")
callout("VAT/세금 별도. 한국 사용자의 경우 결제 시 약 10% 추가될 수 있음.", "note")

doc.add_page_break()

# ════════════ 4단계 ════════════
heading("4단계 — API 키 페이지로 이동", 1)
para("좌측 사이드바에서 [API 키] 메뉴를 클릭합니다.")
para("(좌측 메뉴가 안 보이면 우상단 [조직 설정] 또는 [API 키 받기] 버튼으로 진입)")
para()
para("순서:", bold=True)
para("  ① 좌측 사이드바 [API 키] 클릭 (빨간 박스)")
para("  ② 우상단 [+ 키 생성] 버튼 클릭 (빨간 박스)")

image("4.png", 16)
caption("그림 4. API 키 페이지 — [+ 키 생성] 버튼")

callout("이미 키가 1개 있어도 새 키를 추가로 만들 수 있습니다 (자동입력기용 별도 키 권장).", "note")

doc.add_page_break()

# ════════════ 5단계 ════════════
heading("5단계 — 키 이름 입력 + 추가", 1)
para("[+ 키 생성] 버튼을 누르면 다음 다이얼로그가 나타납니다.")
para()
para("입력:", bold=True)
para("  ● 워크스페이스에서 만들기: [워크스페이스 선택] → 기본값 그대로 또는 Default 선택")
para("  ● 키 이름 지정: 'my-secret-key' 자리에 본인이 알아볼 이름 입력 (예: '문서접수', 'kedu')")
para("  ● [추가] 버튼 클릭")

image("5.png", 10)
caption("그림 5. API 키 생성 다이얼로그")

callout("키 이름은 본인 구분용 라벨일 뿐 — 아무 이름이나 무방.", "info")

doc.add_page_break()

# ════════════ 6단계 ════════════
heading("6단계 — 키 복사 (★ 매우 중요 ★)", 1)
callout("이 화면의 키는 한 번만 표시됩니다. 닫으면 다시 확인 불가 — 반드시 [키 복사] 클릭!", "warn")
para()
para("순서:", bold=True)
para("  ① 화면에 표시된 키 (sk-ant-... 로 시작) 옆 [📋 키 복사] 버튼 클릭")
para("  ② 키가 클립보드에 복사됨 (메모장 등에 임시로 붙여넣어 두면 안전)")
para("  ③ 자동입력기에 등록하기 전까지 [닫기] 누르지 말기")

image("6.png", 10)
caption("그림 6. 생성된 API 키 — [키 복사]")

callout("키 분실 시: 이 페이지를 닫으면 다시 못 봅니다. 새 키를 발급받으세요 (이전 키는 삭제 권장).", "warn")
callout("키 형태: sk-ant-api03-... 로 시작하는 100자 이상의 긴 문자열.", "note")

doc.add_page_break()

# ════════════ 7단계 ════════════
heading("7단계 — 자동입력기에 키 등록", 1)
para("비전자문서등록 자동입력 프로그램(EXE) 을 실행합니다.")
para()
para("순서:", bold=True)
para("  ① GUI 우상단 [⚙ 설정] 버튼 클릭")
para("  ② [Anthropic API 키] 입력칸에 복사한 키 (Ctrl+V) 붙여넣기 — 빨간 박스 위치")
para("  ③ [Claude 모델] 콤보박스에서 원하는 모델 선택 (기본 Sonnet 4.5 권장)")
para("  ④ [저장] 클릭")

image("7.png", 14)
caption("그림 7. 자동입력기 — [⚙ 설정] 다이얼로그에 키 붙여넣기")

callout("키는 본인 PC 의 %USERPROFILE%\\.kedu_anthropic_key 파일에 저장됩니다. 이 파일을 다른 사람에게 공유하지 마세요.", "warn")
callout("모델 추천:\n  • Sonnet 4.5 (기본) — 정확 + 합리적 (~10원/PDF)\n  • Haiku 4.5 — 저렴 (~1원/PDF), 디지털 PDF 위주에 적합\n  • Opus 4.7 — 최고 정확도 (~50원/PDF)", "info")

doc.add_page_break()

# ════════════ FAQ ════════════
heading("자주 묻는 질문 (FAQ)", 1)

heading("Q1. 가입은 무료인가요?", 2)
para("네, Claude Console 가입 자체는 무료입니다. API 호출 비용만 사용량만큼 차감됩니다.")
para("크레딧 충전 ($5 최소) → 사용량만큼 자동 차감.")

heading("Q2. 비용이 얼마나 나오나요?", 2)
para("PDF 1건 처리 평균 비용:")
para("  ● Sonnet 4.5 사용: 약 10원")
para("  ● Haiku 4.5 사용: 약 1원")
para("  ● Opus 4.7 사용: 약 50원")
para("→ $5 충전 (약 7,000원) 시 Sonnet 으로 약 500건 처리 가능.")

heading("Q3. 키를 분실했어요.", 2)
para("키 자체는 다시 볼 수 없습니다. 다음 절차로 새 키 발급:")
para("  ① https://console.anthropic.com/settings/keys 접속")
para("  ② [+ 키 생성] 으로 새 키 발급")
para("  ③ 자동입력기 [⚙ 설정] 에서 새 키 등록")
para("  ④ 분실된 옛 키는 옆의 [⋯] 메뉴에서 [폐기] (보안)")

heading("Q4. 회사 카드 결제 가능한가요?", 2)
para(
    "Anthropic 결제는 Stripe 를 통한 국제 카드 결제입니다. "
    "법인카드/비자/마스터 등 일반 신용카드 모두 가능. VAT 영수증 발급 가능."
)

heading("Q5. 사용량 모니터링은 어디서 보나요?", 2)
para("console.anthropic.com 의 [크레딧] 또는 [분석] 메뉴에서 일/월별 사용량 + 비용 확인.")

heading("Q6. 키가 외부에 노출되면 어떻게 되나요?", 2)
para(
    "키 소지자가 본인 계정으로 API 호출 가능 → 본인에게 청구됨. "
    "GitHub, 채팅, 이메일 등에 절대 공유 금지. "
    "노출 의심 시 즉시 console 에서 키 폐기 후 새 키 발급."
)

doc.add_page_break()

# ════════════ 보안 가이드 ════════════
heading("보안 가이드", 1)
para("Anthropic API 키 보호를 위한 권장 사항:", bold=True)
para()
para("✅ 권장:")
para("  ● 키는 본인 PC 의 %USERPROFILE%\\.kedu_anthropic_key 에만 저장 (자동입력기가 처리)")
para("  ● console.anthropic.com 의 [크레딧] / [분석] 에서 사용량 정기 점검")
para("  ● 30~90일마다 키 갱신 (Rotate)")
para()
para("❌ 금지:")
para("  ● 메일/카카오톡/슬랙 등 메신저에 키 공유")
para("  ● 공용 PC 에서 키 등록 (사용 후 .kedu_anthropic_key 파일 반드시 삭제)")
para("  ● GitHub/Notion 등 외부 시스템에 키 업로드")
para("  ● 같은 키를 여러 사람이 공유 (→ 비용 청구 추적 불가)")

para()
heading("키 갱신(Rotate) 방법", 2)
para("  ① console.anthropic.com/settings/keys 접속")
para("  ② 기존 키 옆 [⋯] → [폐기] 클릭")
para("  ③ [+ 키 생성] 으로 새 키 발급")
para("  ④ 자동입력기 [⚙ 설정] 에서 새 키로 교체")

doc.add_page_break()

heading("부록 — 참고 링크", 1)
para("● Claude Console: https://console.anthropic.com/")
para("● API 키 페이지: https://console.anthropic.com/settings/keys")
para("● 크레딧 / 결제: https://console.anthropic.com/settings/billing")
para("● API 가격표: https://www.anthropic.com/pricing")
para("● 자동입력기 다운로드: https://github.com/xier1108-art/kedu-doc-auto/releases/latest")

para()
para("본 문서는 K-에듀파인 비전자문서등록 자동입력기 v1.0.7 기준입니다 (2026.05).",
     9, color=(120, 120, 120))

doc.save(OUT)
print(f"생성: {OUT}")
print(f"크기: {OUT.stat().st_size / 1024:.1f} KB")

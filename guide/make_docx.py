"""Anthropic API 키 발급 가이드 Word 문서 생성."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm, Inches

DOC_DIR = Path(__file__).parent
IMG_DIR = DOC_DIR / "images"
OUT = DOC_DIR / "Anthropic_API_키_발급_가이드.docx"

doc = Document()

# 페이지 여백
for s in doc.sections:
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)


# ── 헬퍼 ──
def style_default(p):
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(11)


def heading(text: str, level: int = 1, color: tuple = (40, 80, 160)):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(20 - level * 2)
    run.font.color.rgb = RGBColor(*color)
    run.bold = True
    return p


def para(text: str = "", size: int = 11, bold: bool = False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def callout(text: str, kind: str = "info"):
    """강조 박스 (info / warn / note)."""
    icons = {"info": "💡", "warn": "⚠", "note": "📌"}
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icons.get(kind, '•')}  {text}")
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10)
    color = {"info": (40, 80, 160), "warn": (180, 80, 30), "note": (60, 60, 60)}
    run.font.color.rgb = RGBColor(*color.get(kind, (60, 60, 60)))
    run.bold = True


def image(name: str, width_cm: float = 16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(IMG_DIR / name), width=Cm(width_cm))


def caption(text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.italic = True


# ════════════════════════════════════════════════════════
# 제목
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Anthropic API 키 발급 가이드")
run.font.name = "맑은 고딕"
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(40, 80, 160)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("K-에듀파인 비전자문서등록 자동입력기 사용을 위한 사전 준비")
run.font.name = "맑은 고딕"
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 110, 130)

doc.add_paragraph()

# ── 개요 ──
heading("개요", level=1)
para(
    "비전자문서등록 자동입력기는 PDF 의 내용을 자동으로 추출하기 위해 Anthropic Claude API 를 사용합니다. "
    "본 가이드는 Anthropic Console 가입부터 결제 등록, API 키 발급, 그리고 자동입력기에 키를 등록하는 "
    "전 과정을 시각자료와 함께 설명합니다."
)

callout("예상 소요 시간: 약 5-10분 (결제수단 + 신용카드 정보 입력 시간 포함)", kind="info")
callout("최소 결제 금액: $5 (약 7,000원). PDF 1건당 약 10원이므로 500건 이상 사용 가능.", kind="info")
callout("이 키는 본인만 사용. 메신저·이메일·외부 사이트에 절대 공유 금지.", kind="warn")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 1단계: 가입
# ════════════════════════════════════════════════════════
heading("1단계 — Anthropic Console 가입", level=1)
para("브라우저에서 다음 주소로 이동:", size=11)
para("https://console.anthropic.com/", size=12, bold=True, color=(40, 80, 160))

para()
para("가입 방법 (둘 중 하나 선택):", bold=True)
para("  ① Google 계정으로 가입 — 가장 빠름 (Google 로그인 한 번으로 완료)")
para("  ② 이메일로 가입 — name@example.com 형태의 이메일 입력 → 이메일 인증")

image("01_signup.png", 16)
caption("그림 1. Anthropic Console 가입 페이지")

callout("Google 가입을 추천합니다 — 이메일 인증 단계 생략.", kind="info")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 2단계: 결제
# ════════════════════════════════════════════════════════
heading("2단계 — 결제수단 등록 + 크레딧 충전", level=1)
para(
    "Claude API 는 사용량만큼 과금되는 방식입니다 (PDF 1건당 약 10원). "
    "사용 전 결제수단 등록 + 크레딧 충전이 필요합니다."
)
para()
para("순서:", bold=True)
para("  ① 좌측 사이드바에서 [Plans & Billing] 클릭")
para("  ② [Add payment method] 클릭 → 신용/체크카드 정보 입력")
para("  ③ [Buy credits] 클릭 → 최소 $5 충전 (약 7,000원, 500건 분량)")

image("02_billing.png", 16)
caption("그림 2. 결제 페이지 — 카드 등록 + 크레딧 충전")

callout("Auto-reload 기능: 크레딧이 일정 수준 이하로 떨어지면 자동 충전 (선택 사항).", kind="info")
callout("VAT/세금 별도. 한국 사용자의 경우 결제 시 약 10% 추가될 수 있음.", kind="note")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 3단계: API Keys 페이지
# ════════════════════════════════════════════════════════
heading("3단계 — API Keys 페이지로 이동", level=1)
para("좌측 사이드바에서 [API Keys] 메뉴를 클릭합니다.")
para("페이지 우상단의 [+ Create Key] 버튼을 클릭합니다.")

image("03_apikeys.png", 16)
caption("그림 3. API Keys 페이지 — [+ Create Key] 버튼")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 4단계: 키 생성 다이얼로그
# ════════════════════════════════════════════════════════
heading("4단계 — 키 이름 입력 + 생성", level=1)
para("[+ Create Key] 버튼을 누르면 다음 다이얼로그가 나타납니다.")
para()
para("입력:", bold=True)
para("  ① Name: 키 이름을 자유롭게 (예: 'kedu-doc-auto', '문서접수' 등)")
para("  ② Workspace: Default Workspace 그대로 두면 됨")
para("  ③ [Create] 버튼 클릭")

image("04_createkey.png", 16)
caption("그림 4. 키 이름 입력 + Create 클릭")

callout("Name 은 본인이 키를 구분하기 위한 라벨일 뿐 — 아무 이름이나 OK.", kind="info")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 5단계: 키 복사
# ════════════════════════════════════════════════════════
heading("5단계 — 키 복사 (★ 매우 중요 ★)", level=1)
callout("이 화면의 키는 한 번만 표시됩니다. 닫으면 다시 확인할 수 없습니다.", kind="warn")
para()
para("순서:", bold=True)
para("  ① 표시된 키 (sk-ant-... 로 시작) 옆의 [📋 Copy] 버튼 클릭")
para("  ② 키가 클립보드에 복사됨")
para("  ③ [Done] 클릭하기 전에 자동입력기 GUI 로 이동해서 등록 (다음 단계)")

image("05_copykey.png", 16)
caption("그림 5. 생성된 API 키 — [Copy] 로 복사")

callout("키 분실 시: 이 페이지를 닫은 후엔 다시 못 봅니다. 새 키 발급 받으세요 (이전 키는 Revoke).", kind="warn")
callout("키 형태: sk-ant-api03-... 로 시작하는 100자 이상 긴 문자열.", kind="note")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 6단계: 자동입력기에 등록
# ════════════════════════════════════════════════════════
heading("6단계 — 자동입력기에 키 등록", level=1)
para("비전자문서등록 자동입력 프로그램을 실행합니다.")
para()
para("순서:", bold=True)
para("  ① GUI 우상단 [⚙ 설정] 버튼 클릭")
para("  ② [Anthropic API 키] 입력 칸에 복사한 키 (Ctrl+V) 붙여넣기")
para("  ③ Claude 모델 콤보박스에서 원하는 모델 선택 (기본: Sonnet 4.5)")
para("  ④ [저장] 버튼 클릭")

image("06_gui_settings.png", 16)
caption("그림 6. 자동입력기의 [⚙ 설정] 다이얼로그")

callout("키는 본인 PC 의 %USERPROFILE%\\.kedu_anthropic_key 파일에 저장됩니다. 다른 사람에게 PC 의 이 파일을 공유하지 마세요.", kind="warn")
callout("모델 선택:\n  • Sonnet 4.5 (기본 권장) — 정확 + 합리적 비용 (~10원/PDF)\n  • Haiku 4.5 — 저렴 (~1원/PDF), 디지털 PDF 위주면 충분\n  • Opus 4.7 — 최고 정확도 + 비싼 (~50원/PDF)", kind="info")

doc.add_page_break()

# ════════════════════════════════════════════════════════
# FAQ
# ════════════════════════════════════════════════════════
heading("자주 묻는 질문 (FAQ)", level=1)

heading("Q1. 가입은 무료인가요?", level=2)
para("네, Anthropic Console 가입 자체는 무료입니다. API 사용량만큼만 후불 청구됩니다.")
para("크레딧 충전 ($5 최소) 후에 사용량만큼 차감되는 방식.")

heading("Q2. 비용이 얼마나 나오나요?", level=2)
para("• PDF 1건 처리 시 평균 비용:")
para("  - Sonnet 4.5 사용 시: 약 10원")
para("  - Haiku 4.5 사용 시: 약 1원")
para("  - Opus 4.7 사용 시: 약 50원")
para("• $5 충전 시 Sonnet 으로 약 500건 처리 가능 (Haiku 5000건)")

heading("Q3. 키를 분실했어요.", level=2)
para("키 자체는 다시 볼 수 없습니다. 다음 절차로 새 키 발급:")
para("  ① https://console.anthropic.com/settings/keys 접속")
para("  ② 분실한 키 옆의 [Revoke] 클릭 (보안상 권장)")
para("  ③ [+ Create Key] 로 새 키 생성")
para("  ④ 자동입력기 [⚙ 설정] 에서 새 키 등록")

heading("Q4. 회사 카드 결제 가능한가요?", level=2)
para(
    "Anthropic 결제는 Stripe 를 통한 국제 카드 결제입니다. "
    "법인카드/비자/마스터 등 일반 신용카드 모두 가능. VAT 영수증 발급 가능."
)

heading("Q5. 사용량 모니터링은 어디서 보나요?", level=2)
para("console.anthropic.com 좌측 사이드바 [Usage] 메뉴 — 일/월 단위 사용량 + 비용 확인.")

heading("Q6. 키가 외부에 노출되면 어떻게 되나요?", level=2)
para(
    "키를 가진 사람이 본인 계정으로 API 호출 가능 → 본인에게 청구됨. "
    "GitHub, 채팅, 이메일 등에 절대 공유 금지. "
    "노출 의심 시 즉시 console 에서 Revoke 후 새 키 발급."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
# 보안
# ════════════════════════════════════════════════════════
heading("보안 가이드", level=1)
para("Anthropic API 키 보호를 위한 권장 사항:", bold=True)
para()
para("✅ 권장:")
para("  • 키는 본인 PC 의 %USERPROFILE%\\.kedu_anthropic_key 에만 저장 (자동입력기가 처리)")
para("  • 키 사용처를 console.anthropic.com 의 [Usage] 에서 정기적으로 확인")
para("  • 30일 또는 90일마다 키를 갱신 (Rotate)")
para()
para("❌ 금지:")
para("  • 메일/카카오톡/슬랙 등 메신저에 키 공유")
para("  • 공용 PC 에서 키 등록 (사용 후 반드시 .kedu_anthropic_key 삭제)")
para("  • GitHub/Notion 등 외부 시스템에 키 업로드")
para("  • 같은 키를 여러 사람이 공유 (→ 비용 청구 추적 불가)")

para()
heading("키 갱신 (Rotate) 방법", level=2)
para("  ① console.anthropic.com/settings/keys 접속")
para("  ② 기존 키 옆 [Revoke] 클릭")
para("  ③ [+ Create Key] 로 새 키 발급")
para("  ④ 자동입력기 [⚙ 설정] 에서 새 키로 교체")

doc.add_page_break()

heading("부록 — 참고 링크", level=1)
para("• Anthropic Console: https://console.anthropic.com/")
para("• API 키 페이지: https://console.anthropic.com/settings/keys")
para("• 결제 페이지: https://console.anthropic.com/settings/billing")
para("• 사용량 페이지: https://console.anthropic.com/settings/usage")
para("• API 가격표: https://www.anthropic.com/pricing")
para("• 자동입력기 다운로드: https://github.com/xier1108-art/kedu-doc-auto/releases/latest")

para()
para("이 문서는 K-에듀파인 비전자문서등록 자동입력기 v1.0.7 기준입니다.",
     size=9, color=(120, 120, 120))

# 저장
doc.save(OUT)
print(f"생성: {OUT}")
print(f"크기: {OUT.stat().st_size / 1024:.1f} KB")
